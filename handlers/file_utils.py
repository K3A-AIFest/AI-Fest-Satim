"""
File upload utilities for extracting text from various document formats.

This module provides functions to extract text content from uploaded files
including DOCX and PDF formats.
"""
import os
import tempfile
from typing import Union
from pathlib import Path
from fastapi import UploadFile, HTTPException

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text content from an uploaded file.
    
    Args:
        file: The uploaded file (DOCX or PDF)
        
    Returns:
        str: Extracted text content
        
    Raises:
        HTTPException: If file format is not supported or extraction fails
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
        
    file_extension = Path(file.filename).suffix.lower()
    
    if file_extension == '.docx':
        return await _extract_text_from_docx(file)
    elif file_extension == '.pdf':
        return await _extract_text_from_pdf(file)
    else:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file format: {file_extension}. Supported formats: .docx, .pdf"
        )

async def _extract_text_from_docx(file: UploadFile) -> str:
    """Extract text from a DOCX file."""
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500, 
            detail="DOCX support not available. Install python-docx package."
        )
    
    try:
        # Read file content into memory
        content = await file.read()
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        try:
            # Extract text using python-docx
            doc = docx.Document(temp_file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n\n'.join(text_content)
            
        finally:
            # Clean up temporary file
            os.unlink(temp_file_path)
            
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"Error extracting text from DOCX file: {str(e)}"
        )

async def _extract_text_from_pdf(file: UploadFile) -> str:
    """Extract text from a PDF file."""
    if not PDF_AVAILABLE:
        raise HTTPException(
            status_code=500, 
            detail="PDF support not available. Install PyPDF2 package."
        )
    
    try:
        # Read file content into memory
        content = await file.read()
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name
        
        try:
            # Extract text using PyPDF2
            text_content = []
            
            with open(temp_file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
            
            return '\n\n'.join(text_content)
            
        finally:
            # Clean up temporary file
            os.unlink(temp_file_path)
            
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"Error extracting text from PDF file: {str(e)}"
        )

def validate_file_size(file: UploadFile, max_size_mb: int = 10) -> None:
    """
    Validate that the uploaded file is within size limits.
    
    Args:
        file: The uploaded file
        max_size_mb: Maximum file size in MB (default: 10MB)
        
    Raises:
        HTTPException: If file is too large
    """
    if file.size and file.size > max_size_mb * 1024 * 1024:
        raise HTTPException(
            status_code=413, 
            detail=f"File size ({file.size / (1024*1024):.1f}MB) exceeds maximum allowed size ({max_size_mb}MB)"
        )

def validate_file_type(file: UploadFile) -> None:
    """
    Validate that the uploaded file has a supported type.
    
    Args:
        file: The uploaded file
        
    Raises:
        HTTPException: If file type is not supported
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
        
    file_extension = Path(file.filename).suffix.lower()
    supported_extensions = ['.docx', '.pdf']
    
    if file_extension not in supported_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type: {file_extension}. Supported types: {', '.join(supported_extensions)}"
        )
